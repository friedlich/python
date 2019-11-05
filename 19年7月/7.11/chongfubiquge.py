from gevent import monkey
monkey.patch_all()
import gevent, requests, bs4, sys
from gevent.queue import Queue
from urllib.parse import quote
from docx import Document


class Biquge():
    def __init__(self):
        self.headers = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/74.0.3729.169 Safari/537.36'
            }
        # self.proxy = {'http': '117.91.132.241:9999'}
        self.proxy = {}
        self.url_list = []
        self.content_all = []
        self.url = ''
        self.book_name = ''
        self.url_test = []

    def find_book(self):
        self.book_name = input('输入你想查询的小说名称：')
        self.url = 'http://www.biquge.cm/modules/article/sou.php?searchkey=' + quote(self.book_name.encode('gb2312'))
        print(self.url)
        bqg_res = requests.get(self.url, headers=self.headers, proxies=self.proxy)
        bqg_res.encoding = 'GBK'
        bs = bs4.BeautifulSoup(bqg_res.text, 'lxml')
        try_tr = bs.find_all('tr') # 对于大道朝天来说这就是个空列表
        # try_tr = bs.find(class_="grid").find_all('tr') 加了find(class_="grid")就报错，报大道朝天的错
        # try_tr = bs.find(class_="grid").find_all('tr') find(class_="grid")是空的，空对象用不了find_all查找方式
        # AttributeError: 'NoneType' object has no attribute 'find_all' # 因为大道朝天里没有class_="grid"属性，空对象
        while len(try_tr) == 1:
            print('没有在笔趣阁找到该书')
            a = input('输入1继续查询，输入0退出：')
            if a == '1':
                self.book_name = input('输入你想查询的小说名称：')
                self.url = 'http://www.biquge.cm/modules/article/sou.php?searchkey=' + quote(self.book_name.encode('gb2312'))
                print(self.url)
                bqg_res = requests.get(self.url, headers=self.headers, proxies=self.proxy)
                bqg_res.encoding = 'GBK'
                bs = bs4.BeautifulSoup(bqg_res.text, 'lxml')
                # try_tr = bs.find(class_="grid").find_all('tr') 加了find(class_="grid")就报错
                try_tr = bs.find_all('tr') # 记住while是一个循环，就是说如果这里又输了一个错误名称，len(try_tr)又是1，满足while循环不会出来，输对名称才能跳出循环
            elif a == '0':
                return '0'
        if try_tr == []:
            name = bs.find(id="info").find('h1').text
            author = bs.find(id="info").find('p').text.split('：')[1]
            self.book_name = name
            print('已经匹配到合适结果：', '书名：', name,'作者：', author)
            x = input('输入0退出，输入1开始下载本书：')
            return x
            return self.url
        else:
            i = 1
            print('在笔趣阁找到以下相似书名：')
            list_choose = []
            for tr in try_tr[1:]:
                name = tr.find_all(class_="odd")[0].text
                author = tr.find_all(class_="odd")[1].text
                url_choose = tr.find('a')['href']
                list_choose.append({'name': name, 'author': author, 'url': url_choose})
                print(i, name, '作者：', author)
                i += 1
            x = input('输入0退出，输入序号开始下载对应书籍：')
            self.url = list_choose[int(x)-1]['url']
            return x
            return self.url
        
    def get_url(self):
        res = requests.get(self.url, headers=self.headers, proxies=self.proxy)
        res.encoding = 'GBK'
        bs = bs4.BeautifulSoup(res.text, 'lxml')
        # 提取章节标题和内容
        dds = bs.find(id="list").find_all('dd')
        for dd in dds:
            url = 'http://www.biquge.cm' + dd.find('a')['href']
            self.url_list.append(url)
        return self.url_list

    def get_content_multi(self, thread):
        def get_():
            while not work.empty():
                url = work.get_nowait()
                res = requests.get(url, headers=self.headers, proxies=self.proxy)
                res.encoding = 'GBK'
                bs = bs4.BeautifulSoup(res.text, 'lxml')
                title = bs.find(class_="bookname").find('h1').text
                content = bs.find(id="content").text
                content_1 = {}
                num = url[url.rfind('/', 1)+1:-5]
                content_1['num'] = num
                content_1['title'] = title
                content_1['content'] = content.replace('（未完待续）', '')
                content_all.append(content_1)
                print(content_1['title'], '目前还剩', work.qsize(), '章')               
                return content_all
        work = Queue()
        content_all = []
        for i in range(5):
            self.url_test.append(self.url_list[i])
        print(self.url_test)
        for url in self.url_test:
            work.put_nowait(url)
        task_list = []
        for x in range(thread):
            task = gevent.spawn(get_)
            task_list.append(task)
        gevent.joinall(task_list)
        self.content_all = content_all
        self.content_all = sorted(self.content_all, key=lambda keys: keys['num'])

    def save_txt(self):
        with open(sys.path[0]+'\\'+self.book_name+'.txt', 'w', encoding='utf-8') as f:
            for x in self.content_all:
                f.write(x['title']+'\n')
                f.write(x['content'])

    def save_docx(self):
        # doc = Document(self.book_name+'docx')
        doc = Document(sys.path[0]+'\\'+self.book_name+'.docx')
        for x in self.content_all:
            doc.add_heading(x['title'], level=2)
            doc.add_paragraph(x['content'])
        # doc.save(self.book_name+'docx') # 又是一个报错
        # doc.save(sys.path[0]+'\\大道朝天.docx')
        doc.save(sys.path[0]+'\\'+self.book_name+'.docx')


if __name__ == '__main__':
    print(__name__)
    download = Biquge()
    x = download.find_book()
    if x != '0':
        download.get_url()
        download.get_content_multi(3)
        # download.save_txt()
        download.save_docx()



