from gevent import monkey
monkey.patch_all()
import gevent, requests, csv
from gevent.queue import Queue
from bs4 import BeautifulSoup
from docx import Document
import urllib

class Biquge():
    def __init__(self):
        self.headers={
            'User - Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36 SE 2.X MetaSr 1.0'
            }
        self.url_list = []
        self.content_all = []
        self.url = ''
        self.book_name = ''

    def find_book(self):
        self.book_name = input('输入你想查询的小说名称')
        self.url = 'http://www.biquge.cm/modules/article/sou.php?searchkey=' + urllib.parse.quote(
            self.book_name.encode('gb2312'))
        print(self.url)
        bqg_res = requests.get(self.url, headers=self.headers)
        bqg_res.encoding = 'GBK'
        bqg_res = bqg_res.text
        bs=BeautifulSoup(bqg_res,'html.parser')
        try_tr=bs.find_all('tr')
        while len(try_tr)==1:
            print('没有在笔趣阁找到该书')
            a=input('输入1继续查询，输入0退出')
            if a=='1':
                self.book_name = input('输入你想查询的小说名称')
                self.url = 'http://www.biquge.cm/modules/article/sou.php?searchkey=' + urllib.parse.quote(
                    self.book_name.encode('gb2312'))
                bqg_res = requests.get(self.url, headers=self.headers)
                bqg_res.encoding = 'GBK'
                bqg_res = bqg_res.text
                bs = BeautifulSoup(bqg_res, 'html.parser')
                try_tr = bs.find_all('tr')
            elif a=='0':
                return '0'

        if try_tr==[]:
            name=bs.find('div',id='info').find('h1').text
            author = bs.find('div', id='info').find('p').text.replace('    ','')
            self.book_name=name
            print('已经匹配到合适结果', '书名：', name, author)
            x=input('输入0退出，输入1开始下载本书')
            return  x
        else:
            i = 1
            print('在笔趣阁找到以下相似书名')
            list_choose=[]
            for tr in try_tr[1:]:
                name = tr.find('a').text
                author = tr.find_all('td', class_='odd')[1].text
                url_choose=tr.find('a')['href']
                list_choose.append({'name':name,'author':author,'url':url_choose})
                print(i, name, ' 作者：', author)
                i += 1
            x = input('输入0退出，输入序号开始下载对应书籍')
            self.url=list_choose[int(x)-1]['url']
            return x

    def get_url(self):
        res = requests.get(self.url,headers=self.headers)
        res.encoding = 'GBK'
        res = res.text
        bs = BeautifulSoup(res, 'html.parser')
        dds = bs.find('div', id='list').find_all('dd')
        for dd in dds:
            url = 'http://www.biquge.cm' + dd.find('a')['href']
            self.url_list.append(url)
        return self.url_list

    def get_content_multi(self,thread):
        def get_():
            while not work.empty():
                url = work.get_nowait()
                res = requests.get(url, headers=self.headers)
                res.encoding = 'GBK'
                res = res.text
                bs = BeautifulSoup(res, 'html.parser')
                # 提取章节标题和内容
                title = bs.find('div', class_='bookname').find('h1').text
                content = bs.find('div', id='content').text
                # 构建字典存储，方便后面排序
                content_1 = {}
                num = url[url.rfind('/', 1) + 1:-5]
                content_1['num'] = num
                content_1['title'] = title
                content_1['content'] = content.replace('（未完待续）', '')
                content_all.append(content_1)
                print(content_1['title'], '目前还剩', work.qsize(), '章')

        work = Queue()
        content_all=[]
        for url in self.url_list:
            work.put_nowait(url)
        tasks_list = []
        for x in range(thread):
            task = gevent.spawn(get_)
            tasks_list.append(task)
        gevent.joinall(tasks_list)
        self.content_all=content_all
        self.content_all = sorted(self.content_all, key=lambda keys: keys['num'])

    def save_docx(self):
        doc = Document(self.book_name+'docx')
        for x in self.content_all:
            doc.add_heading(x['title'], level=2)
            doc.add_paragraph(x['content'])
        doc.save(self.book_name+'docx')

    def save_txt(self):
        with open(self.book_name+'txt','w',encoding='utf-8') as f:
            for x in self.content_all:
                f.write(x['title']+'\n')
                f.write(x['content'])


if __name__ == '__main__':
    download=Biquge()
    x=download.find_book()
    if x!='0':
        download.get_url()
        download.get_content_multi(2)
        download.save_txt()







