from gevent import monkey
monkey.patch_all()
import gevent,requests,csv
from gevent.queue import Queue
from bs4 import BeautifulSoup
from docx import Document

headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36(KHTML, like Gecko) Chrome/66.0.3359.139 Safari/537.36'}

def get_url():
    url_0 = 'http://www.biquge.cm/3/3095/'
    res = requests.get(url_0,headers=headers)
    res.encoding = 'GBK'
    bs = BeautifulSoup(res.text,'html.parser')
    dds = bs.find(id="list").find_all('dd')
    for dd in dds:
        url = 'http://www.biquge.cm' + dd.find('a')['href']
        url_list.append(url)

def get_content():
    while not work.empty(): #当队列不是空的时候，就执行下面的程序
        url = work.get_nowait()
        #用get_nowait()函数可以把队列里的网址都取出
        res = requests.get(url,headers=headers)
        res.encoding = 'GBK'
        bs = BeautifulSoup(res.text,'html.parser')
        # 提取章节标题和内容
        title = bs.find(class_="bookname").find('h1').text
        content = bs.find(id="content").text
        # 构建字典存储，方便后面排序
        content_1 = {}
        num = url.replace('http://www.biquge.cm/3/3095/','').replace('.html','')
        content_1['num'] = num
        content_1['title'] = title
        content_1['content'] = content.replace('（未完待续）','')
        content_all.append(content_1)
        print(content_1['title'], '目前还剩', work.qsize(), '章') # qsize方法，是用来判断队列里还剩多少数量的

def save_content(book):
    doc = Document(book)
    for x in content_all:
        doc.add_heading(x['title'], level=2)
        doc.add_paragraph(x['content'])
        doc.save(book)

if __name__ == '__main__':
    work = Queue()
    #创建队列对象，并赋值给work
    url_list = []
    url_test = []
    tasks_list = []
    content_all = [] 

    get_url()

    for i in range(5):
        url_test.append(url_list[i])
    print(url_test)

    for x in url_test:
        work.put_nowait(x)
        #用put_nowait()函数可以把网址都放进队列里

    for x in range(3): #相当于创建了3个爬虫
        task = gevent.spawn(get_content)
        #用gevent.spawn()函数创建执行get_content(content_all)函数的任务
        tasks_list.append(task)
        #往任务列表添加任务
    gevent.joinall(tasks_list)
    #用gevent.joinall方法，执行任务列表里的所有任务，就是让爬虫开始爬取网站

    content_all = sorted(content_all,key = lambda keys:keys['num'])

    save_content(r'c:\Users\asus\Desktop\Python\风变Python爬虫精进\爬虫阶段练习\19年7月\7.10\妙手天师.docx')




