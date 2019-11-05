from gevent import monkey
monkey.patch_all()
import gevent, requests, csv
from gevent.queue import Queue
from bs4 import BeautifulSoup
from docx import Document
# Python可以利用python-docx模块处理word文档，处理方式是面向对象的。也就是说python-docx模块会把word文档，
# 文档中的段落、文本、字体等都看做对象，对对象进行处理就是对word文档的内容处理。

def get_content(content_all):
    while not work.empty():
        url = work.get_nowait()
        res = requests.get(url)
        res.encoding = 'GBK'
        res = res.text
        bs = BeautifulSoup(res, 'html.parser')
        # 提取章节标题和内容
        title = bs.find('div', class_='bookname').find('h1').text
        content = bs.find('div', id='content').text
        # 构建字典存储，方便后面排序
        content_1 = {}
        num = url.replace('http://www.biquge.cm/3/3095/', '').replace('.html', '')
        content_1['num'] = num
        content_1['title'] = title
        content_1['content'] = content.replace('（未完待续）','')
        content_all.append(content_1)
        print(content_1['title'],'目前还剩', work.qsize(), '章')

def save_content(book):
    doc = Document(book)
    for x in content_all:
        doc.add_heading(x['title'], level=2)
        doc.add_paragraph(x['content'])
        doc.save(book)

def get_url(url_list):
    url_0='http://www.biquge.cm/3/3095'
    res = requests.get(url_0)
    res.encoding = 'GBK'
    res = res.text
    bs = BeautifulSoup(res, 'html.parser')
    dds = bs.find('div',id='list').find_all('dd')
    for dd in dds:
        url='http://www.biquge.cm'+dd.find('a')['href']
        url_list.append(url)

if __name__ == '__main__':
    work = Queue()
    url_list = []
    tasks_list = []
    content_all=[]

    get_url(url_list)

    for x in url_list:
        work.put_nowait(x)

    for x in range(2):
        task = gevent.spawn(get_content,content_all)
        tasks_list.append(task)
        gevent.joinall(tasks_list)

    content_all=sorted(content_all,key = lambda keys:keys['num'])

    save_content(r'c:\Users\asus\Desktop\Python\风变Python爬虫精进\爬虫阶段练习\19年7月\7.10\妙手天师.docx')

    


    




